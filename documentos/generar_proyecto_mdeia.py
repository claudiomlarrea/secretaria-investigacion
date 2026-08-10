#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera Anexo I — Proyecto MDeIA UCCuyo (Observatorio de IA)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "Anexo-I-Proyecto-MDeIA-OIA-UCCuyo.docx"

TITULO = (
    "Modelo de Madurez Digital e Inteligencia Artificial (MDeIA): "
    "línea de base del IMD, gobernanza de IA y gemelo digital del plan institucional en la UCCuyo"
)
TITULO_CORTO = "MDeIA UCCuyo — IMD y gobernanza de IA"
CONVOCATORIA = "Secretaría de Investigación · UCCuyo"
FECHA = date.today().strftime("%d/%m/%Y")


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, *, bold: bool = False) -> None:
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> Document:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("UNIVERSIDAD CATÓLICA DE CUYO\n")
    r.bold = True
    r.font.size = Pt(12)
    r2 = title.add_run("Secretaría de Investigación y Vinculación Tecnológica\n\n")
    r2.font.size = Pt(11)
    r3 = title.add_run("PRESENTACIÓN DE PROYECTO DE INVESTIGACIÓN")
    r3.bold = True
    r3.font.size = Pt(14)

    doc.add_paragraph()
    _p(doc, f"Título del Proyecto: {TITULO}", bold=True)
    _p(doc, "Director/a: José La Malfa")
    _p(doc, "Co-director/a: Claudio Marcelo Larrea")
    _p(doc, "Unidad Académica: Observatorio de Inteligencia Artificial")
    _p(doc, f"Convocatoria / Año: {CONVOCATORIA} / 2026")
    _p(doc, f"Fecha de presentación: {FECHA}")

    _h(doc, "1. Identificación del proyecto")
    _p(doc, f"Denominación del Proyecto: {TITULO}")
    _p(doc, f"Denominación abreviada: {TITULO_CORTO}")
    _p(doc, "Tipo de Proyecto: Aplicado / Innovación (investigación-acción institucional)")
    _p(doc, "Duración estimada (meses): 24")
    _p(doc, "Unidad Académica / Instituto / Centro: Observatorio de Inteligencia Artificial — Secretaría de Investigación")
    _p(doc, "Área disciplinar: Ciencias de la computación; Educación; Gestión universitaria; Interdisciplinario")

    _h(doc, "2. Datos del Director/a y Co-director/a")
    _p(doc, "Director/a: José La Malfa")
    _p(doc, "Cargo académico: [Completar — ej. Decano / Director de unidad / Prof. Titular]")
    _p(doc, "Formación académica: [Completar — título de grado y posgrado máximo]")
    _p(doc, "Categoría de investigador/a: [Completar según registro institucional]")
    _p(doc, "Correo institucional: [Completar @uccuyo.edu.ar]")
    doc.add_paragraph()
    _p(doc, "Co-director/a: Claudio Marcelo Larrea")
    _p(doc, "Cargo académico: Responsable del Observatorio de Inteligencia Artificial")
    _p(doc, "Formación académica: [Completar]")
    _p(doc, "Categoría de investigador/a: [Completar]")
    _p(doc, "Correo institucional: observatorioia@uccuyo.edu.ar")

    _h(doc, "3. Equipo de investigación")
    _bullets(
        doc,
        [
            "Belén Arias — Observatorio de IA / UCCuyo — Diseño metodológico, revisión de instrumentos y análisis cualitativo.",
            "Javier Coria — Observatorio de IA / UCCuyo — Ingeniería de datos, tableros, integración con planillas institucionales.",
            "Stefania Young — Observatorio de IA / UCCuyo — Coordinación de campo, encuestas y articulación con unidades académicas.",
            "Laura Pizarro — Observatorio de IA / UCCuyo — Comunicación científica, difusión y transferencia de resultados.",
        ],
    )

    _h(doc, "4. Resumen del proyecto (≈300 palabras)")
    _p(
        doc,
        "La Universidad Católica de Cuyo atraviesa un proceso de transformación digital acelerado "
        "por el uso de inteligencia artificial en docencia, investigación, gestión y extensión. "
        "Sin embargo, carece de un diagnóstico sistemático, comparable y reiterable de su madurez "
        "digital e IA a nivel institucional, por sede y por unidad académica. Este proyecto propone "
        "implementar y validar el Modelo de Madurez Digital e Inteligencia Artificial (MDeIA UCCuyo), "
        "adaptación del marco UDigital madurez (MetaRed) con extensión propia del Observatorio de IA, "
        "mediante el Índice de Madurez Digital (IMD) y un conjunto de 129 indicadores (36 en línea de base). "
        "El estudio combina autodiagnóstico participativo, carga estructurada de evidencias, encuesta "
        "estudiantil sobre IA y construcción de un gemelo digital del Plan Estratégico Institucional "
        "para apoyar la toma de decisiones. La unidad de análisis incluye la institución completa, "
        "las sedes Mendoza, San Luis y San Juan, y cada facultad/unidad transversal. Se espera obtener "
        "una línea de base 2026 del IMD, brechas prioritarias, recomendaciones de gobernanza de IA "
        "(guía ética, políticas de uso en evaluaciones) y herramientas de software abiertas para "
        "monitoreo periódico. El enfoque es mixto, con diseño investigación-acción institucional.",
    )

    _h(doc, "5. Fundamentación (400–600 palabras)")
    _p(
        doc,
        "Las universidades latinoamericanas incorporan IA generativa sin marcos homogéneos de "
        "evaluación de madurez digital. MetaRed promueve UDigital madurez como referente regional; "
        "la UCCuyo requiere extender ese marco hacia dimensiones de IA responsable, alineadas a "
        "su identidad católica y regional. Antecedentes empíricos muestran heterogeneidad entre sedes "
        "y unidades en infraestructura TI, formación docente y políticas de evaluación auténtica. "
        "El Observatorio de IA ya desarrolló instrumentos piloto (encuesta estudiantil, indicadores "
        "MDeIA_IA_*), pero falta consolidación metodológica, validación y escalamiento institucional. "
        "La brecha de conocimiento consiste en integrar madurez digital clásica con gobernanza de IA "
        "en un modelo único, operable por no especialistas y trazable al PEI. El marco conceptual "
        "articula: (a) madurez digital por áreas UDigital; (b) dimensión IA (ética, competencias, "
        "infraestructura, datos); (c) agregación multi-nivel (facultad, sede, institución). "
        "El valor agregado es metodológico e institucional: un sistema vivo de medición — no un "
        "informe estático — reutilizable en ciclos de mejora continua y vinculado al Plan "
        "Estratégico de I+D de la universidad.",
    )

    _h(doc, "6. Pertinencia y relevancia (200–400 palabras)")
    _p(doc, "Alineación con el PEI — Objetivo 1 (Sistema integral de aseguramiento de la calidad): "
         "el MDeIA aporta indicadores verificables de procesos digitales e IA en docencia e investigación.")
    _p(doc, "Objetivo 2 (Integralidad, vinculación y comunicación): fortalece comunicación institucional "
         "y vinculación tecnológica mediante tableros y gemelo digital del plan.")
    _p(doc, "Objetivo 3 (Educación a distancia): evalúa madurez de plataformas, LMS y servicios digitales.")
    _p(doc, "Plan Estratégico de la Función I+D: contribuye a objetivos de fortalecimiento de capacidades "
         "de investigación aplicada, transferencia y alianzas (p. ej. objetivos de articulación inter-sede "
         "y generación de evidencia para políticas institucionales).")
    _p(doc, "Aporte institucional: mejora prácticas académicas mediante autodiagnóstico por unidad; "
         "fortalece formación en IA responsable; beneficia a estudiantes (encuesta OIA), docentes "
         "(guías de uso) y autoridades (IMD por sede). Línea prioritaria del Observatorio de IA: "
         "investigación, formación y vinculación sobre impacto ético y regional de la IA.")

    _h(doc, "7. Planteo del problema y objetivos")
    _p(doc, "Planteo del problema: La UCCuyo no dispone de una medición integrada, periódica y "
         "desagregada de su madurez digital e IA, lo que dificulta priorizar inversiones, "
         "políticas de evaluación y gobernanza de IA generativa en las tres sedes.")
    _p(doc, "Objetivo general: Implementar y validar el modelo MDeIA UCCuyo para establecer la "
         "línea de base del IMD institucional, por sede y por unidad académica, y generar "
         "recomendaciones de política y gobernanza de IA.")
    _bullets(
        doc,
        [
            "OE1. Operacionalizar el catálogo de indicadores (129; línea de base 36) y protocolos de autodiagnóstico.",
            "OE2. Obtener la línea de base 2026 del IMD en al menos el 80 % de las unidades académicas definidas.",
            "OE3. Caracterizar percepciones y prácticas estudiantiles sobre IA (encuesta OIA) e integrarlas al modelo.",
            "OE4. Desarrollar el gemelo digital del PEI/plan I+D como herramienta de apoyo a la gestión.",
            "OE5. Elaborar informes de brechas, recomendaciones y plan de mejora 2026–2028.",
        ],
    )

    _h(doc, "8. Originalidad y aporte al conocimiento (100–200 palabras)")
    _p(
        doc,
        "El proyecto aporta un modelo híbrido UDigital + IA con agregación multi-nivel (pooled por sede) "
        "aplicado a una universidad multicampus. Supera diagnósticos ad hoc al estandarizar indicadores, "
        "software de autodiagnóstico y trazabilidad PEI. Innovación pedagógica: encuesta estudiantil "
        "vinculada a indicadores institucionales. Innovación tecnológica: gemelo digital y tableros "
        "Streamlit/Looker integrados a fuentes institucionales. Impacto potencial: replicabilidad en "
        "universidades regionales y publicaciones sobre gobernanza de IA en educación superior.",
    )

    _h(doc, "9. Marco teórico y estado del arte (600–800 palabras)")
    _p(
        doc,
        "El marco se sustenta en modelos de madurez digital organizacional (CMMI, COBIT adaptados "
        "al sector público y educativo), el programa UDigital madurez de MetaRed, y literatura reciente "
        "sobre IA en educación superior (UNESCO, 2023; EU AI Act; informes 2024–2025 sobre "
        "literacy y políticas de integración curricular). Conceptos clave: madurez digital, "
        "transformación digital, IA generativa, gobernanza algorítmica, evaluación auténtica, "
        "investigación-acción institucional. Estudios recientes evidencian brechas entre política "
        "declarada y práctica docente en LATAM; pocos modelos integran IA en el índice de madurez. "
        "A nivel nacional, universidades implementan planes TI aislados; la UCCuyo avanza con "
        "Observatorio de IA pero sin índice consolidado. La comparación internacional (MetaRed, "
        "España, México) muestra utilidad de indicadores comunes inter-institucionales. "
        "Este proyecto posiciona a la UCCuyo como referente regional en medición MDeIA.",
    )

    _h(doc, "10. Metodología (800–900 palabras)")
    _p(doc, "Enfoque: mixto, con predominancia cuantitativa en el IMD y componente cualitativo "
         "en talleres de validación y análisis de brechas.")
    _p(doc, "Diseño: investigación-acción institucional en cuatro fases — (F1) diseño/validación "
         "de instrumentos; (F2) línea de base IMD; (F3) encuesta estudiantil IA; (F4) informes "
         "y transferencia.")
    _p(doc, "Población y muestra: universo de unidades académicas UCCuyo (sedes Mendoza, San Luis, "
         "San Juan; unidades transversales). Muestra estudiantil: estratificada por sede y carrera "
         "en encuesta OIA (meta n ≥ 400 respuestas válidas).")
    _p(doc, "Instrumentos: catálogo MDeIA (129 indicadores; 36 piloto), plantilla Excel, aplicación "
         "web de autodiagnóstico, encuesta estudiantil, rúbricas de satisfacción por indicador, "
         "guías de sesión de 90 minutos por unidad.")
    _p(doc, "Validación: juicio de expertos (Consejo de Investigación, TI, OIA); pilotaje en 2–3 "
         "unidades; revisión de consistencia problema–objetivos–indicadores.")
    _p(doc, "Análisis: IMD = (indicadores satisfechos / evaluados) × 100; agregación pooled por sede; "
         "desgloses por área UDigital y dimensión IA; análisis de brechas prioritarias; "
         "estadística descriptiva en encuesta; visualización en tableros.")
    _p(doc, "Ética: consentimiento informado en encuesta; anonimización; datos agregados en informes; "
         "aprobación según normativa institucional de investigación con humanos si corresponde.")

    _h(doc, "11. Factibilidad y cronograma")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Etapa / Actividad"
    hdr[1].text = "Período estimado"
    hdr[2].text = "Responsables"
    rows = [
        ("Ajuste instrumentos MDeIA y validación", "Meses 1–3", "Larrea, Arias, Coria"),
        ("Línea de base IMD por unidad académica", "Meses 4–12", "Equipo + unidades"),
        ("Encuesta estudiantil IA (OIA)", "Meses 6–10", "Young, Pizarro"),
        ("Gemelo digital PEI / tableros", "Meses 8–14", "Coria, Larrea"),
        ("Informes por sede e institucional", "Meses 13–18", "Larrea, Arias"),
        ("Difusión y transferencia", "Meses 19–24", "Pizarro, equipo"),
    ]
    for act, per, resp in rows:
        row = table.add_row().cells
        row[0].text = act
        row[1].text = per
        row[2].text = resp
    _p(doc, "Recursos: infraestructura del Observatorio de IA, software libre (Python, Streamlit), "
         "planillas institucionales existentes. Riesgos: baja participación de unidades — mitigación "
         "con acompañamiento y Secretaría de Investigación.")

    _h(doc, "12. Impacto esperado y plan de difusión / transferencia")
    _p(doc, "Impacto académico: artículos en congresos (MetaRed, educación superior), informe técnico "
         "MDeIA, dataset anonimizado de indicadores.")
    _p(doc, "Impacto institucional: línea de base IMD 2026, insumos para PEI y plan I+D, guía de "
         "gobernanza de IA.")
    _p(doc, "Difusión: jornadas por sede, sitio del Observatorio, informes HTML/PDF, repositorio institucional.")
    _p(doc, "Transferencia: aplicación MDeIA UCCuyo, plantillas Excel, manual de autodiagnóstico, tableros.")

    _h(doc, "13. Presupuesto, sostenibilidad y alineación institucional")
    _p(doc, "Presupuesto total estimado: $ 0 (recursos humanos e infraestructura institucional existentes).")
    _p(doc, "Fuentes de financiamiento: internas — Observatorio de IA / Secretaría de Investigación.")
    _p(doc, "Sostenibilidad: modelo re-ejecutable anualmente; mantenimiento en repositorio institucional; "
         "integración con Consejo de Investigación y producción científica.")
    _p(doc, "Alineación: PEI (OG calidad, vinculación, EaD); Plan Estratégico I+D UCCuyo; Ordenanza de investigación.")

    _h(doc, "14. Bibliografía (APA 7 — selección)")
    _bullets(
        doc,
        [
            "MetaRed. (s.f.). UDigital madurez. Red Iberoamericana de Innovación y Conocimiento.",
            "UNESCO. (2023). Guidance for generative AI in education and research.",
            "European Commission. (2024). AI Act — regulatory framework.",
            "Salmona, M., & Lievano, R. (2015). Metodología de la investigación (6.ª ed.). Alfaomega.",
            "Observatorio de Inteligencia Artificial UCCuyo. (2025–2026). Documentación MDeIA UCCuyo.",
        ],
    )

    doc.add_paragraph()
    _p(doc, "Firma del Director/a")
    _p(doc, "Firma y aclaración: José La Malfa")
    _p(doc, "Unidad Académica: Observatorio de Inteligencia Artificial")
    _p(doc, f"Fecha: {FECHA}")

    return doc


if __name__ == "__main__":
    build().save(OUT)
    print(f"Generado: {OUT}")
